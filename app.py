import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import seaborn as sns
from docx import Document
from docx.shared import Inches
from io import BytesIO
import base64

# --- MAIN PAGE HEADERS ---
st.title("ğŸŒ€ Ultra Digital Twin: Attock Spiral Concentrator")
st.markdown("### Process Optimization, Economics & KPI Performance Dashboard")
st.write("---")

# --- 1. SETTINGS & STYLING ---
st.set_page_config(page_title="Ultra Spiral Digital Twin", layout="wide")

st.markdown("""
Â  Â  <style>
Â  Â  .stApp { background-color: #0e1117; color: #ffffff; }
Â  Â  [data-testid="stSidebar"] { background-color: #1e2a38; color: white; }
Â  Â  [data-testid="stSidebar"] .stSlider label { color: #ffcc00 !important; }
Â  Â  [data-testid="stMetricValue"] { color: #00ffcc !important; font-size: 30px; }
Â  Â  h1, h2, h3 { color: #ff4b4b; font-family: 'Arial'; }
Â  Â  </style>
Â  Â  """, unsafe_allow_html=True)

# --- 2. BASE DATA ---
FEED_GRADES = {"Gold": 0.80, "Magnetite": 6.0, "Ilmenite": 1.5, "Rutile": 0.30, "Monazite": 0.15}
BASE_REC = {"Gold": 65.0, "Magnetite": 70.0, "Ilmenite": 60.0, "Rutile": 55.0, "Monazite": 50.0}

# --- 3. SIDEBAR CONTROLS ---
st.sidebar.header("ğŸ•¹ï¸ Global Feed Controls")
f_rate = st.sidebar.slider("Feed Rate (tph)", 100, 500, 300)
solids = st.sidebar.slider("Solids %", 10, 50, 30)
d80 = st.sidebar.slider("Feed d80 (microns)", 20, 1000, 150)

st.sidebar.markdown("---")
st.sidebar.subheader("ğŸŒ€ Spiral Mechanics")
splitter_pos = st.sidebar.slider("Splitter Position (0=Tight, 2=Wide)", 0.0, 2.0, 1.0)

st.sidebar.markdown("---")
st.sidebar.header("ğŸ¯ Recovery Targets (%)")
user_targets = {m: st.sidebar.slider(f"{m} Recovery", 0.0, 100.0, float(v)) for m, v in BASE_REC.items()}

st.sidebar.header("ğŸ¯ Target Grade Settings")
user_grade_targets = {m: st.sidebar.slider(f"{m} Target Grade", 0.0, 50.0, 5.0) for m in FEED_GRADES.keys()}

st.sidebar.header("ğŸ’° Economic Pricing ($)")
user_prices = {
Â  Â  "Gold": st.sidebar.number_input("Gold ($/g)", value=80.0),
Â  Â  "Magnetite": st.sidebar.number_input("Magnetite ($/t)", value=100.0),
Â  Â  "Ilmenite": st.sidebar.number_input("Ilmenite ($/t)", value=250.0),
Â  Â  "Rutile": st.sidebar.number_input("Rutile ($/t)", value=800.0),
Â  Â  "Monazite": st.sidebar.number_input("Monazite ($/t)", value=1500.0)
}

# ----------- ADDITION: ECONOMICS / OPEX (NO CHANGE ABOVE) -----------
st.sidebar.markdown("---")
st.sidebar.header("ğŸ’¼ Operating Cost (OPEX)")

labour_cost = st.sidebar.number_input("Labour Cost ($/hr)", value=120.0)
power_kw = st.sidebar.number_input("Installed Power (kW)", value=150.0)
power_rate = st.sidebar.number_input("Electricity Cost ($/kWh)", value=0.12)
water_cost = st.sidebar.number_input("Water Cost ($/hr)", value=15.0)
maintenance_cost = st.sidebar.number_input("Maintenance Cost ($/hr)", value=40.0)
mining_cost = st.sidebar.number_input("Mining Cost ($/ton)", value=8.0)
lease_tax = st.sidebar.number_input("Lease / Rent / Taxes ($/hr)", value=25.0)

st.sidebar.markdown("### ğŸ“ˆ Margin Control")
target_margin = st.sidebar.slider("Target Profit Margin (%)", 0, 60, 25)

# ----------- ADDITION: KPI EVALUATION -----------



# ----------- ADDITION: KPI TARGETS (SIDEBAR ONLY) -----------
st.sidebar.markdown("---")
st.sidebar.header("ğŸ“Œ KPI Target Controls")

target_throughput = st.sidebar.slider(
Â  Â  "Target Throughput (tph)", 100, 500, f_rate
)


target_profit_hr = st.sidebar.number_input(
Â  Â  "Target Profit ($/hr)", value=5000.0
)


# --- 4. ENGINE LOGIC ---
def ultra_engine(rate, targets, prices, split_pos, size_d80):
Â  Â  size_factor = 1.0
Â  Â  if size_d80 < 100: size_factor = size_d80 / 100
Â  Â  elif size_d80 > 400: size_factor = max(0.4, 1.0 - (size_d80 - 400) / 800)
Â  Â Â 
Â  Â  mass_pull = 0.10 + (split_pos * 0.10)
Â  Â  conc_mass = rate * mass_pull
Â  Â Â 
Â  Â  results = []
Â  Â  total_rev = 0
Â  Â  for m, rec in targets.items():
Â  Â  Â  Â  eff_rec = rec * size_factor
Â  Â  Â  Â  feed_m = rate * (FEED_GRADES[m] if m == "Gold" else FEED_GRADES[m]/100)
Â  Â  Â  Â  conc_m = feed_m * (eff_rec/100)
Â  Â  Â  Â  grade = conc_m / conc_mass if m == "Gold" else (conc_m/conc_mass)*100
Â  Â  Â  Â  rev = conc_m * prices[m]
Â  Â  Â  Â  total_rev += rev
Â  Â  Â  Â  results.append({
Â  Â  Â  Â  Â  Â  "Mineral": m,
Â  Â  Â  Â  Â  Â  "Recovery %": round(eff_rec, 2),
Â  Â  Â  Â  Â  Â  "Conc Grade": round(grade, 4),
Â  Â  Â  Â  Â  Â  "Revenue $/hr": round(rev, 2)
Â  Â  Â  Â  })
Â  Â  return pd.DataFrame(results), conc_mass, total_rev

df_res, c_mass, total_revenue = ultra_engine(f_rate, user_targets, user_prices, splitter_pos, d80)
def generate_heatmap_data(targets, prices, d80, solids):
Â  Â  # Variables for heatmap axis
Â  Â  rates = np.linspace(100, 500, 10)Â  # Feed Rate range
Â  Â  splits = np.linspace(0, 2, 10)Â  Â  Â # Splitter range
Â  Â  grid = np.zeros((10, 10))

Â  Â  for i, r in enumerate(rates):
Â  Â  Â  Â  for j, s in enumerate(splits):
Â  Â  Â  Â  Â  Â  # Engine ko call karke profit nikalna
Â  Â  Â  Â  Â  Â  _, _, total_rev = ultra_engine(r, targets, prices, s, d80)
Â  Â  Â  Â  Â  Â Â 
Â  Â  Â  Â  Â  Â  # OPEX Calculate karna usi logic se jo aapne niche likhi hai
Â  Â  Â  Â  Â  Â  p_cost = power_kw * power_rate
Â  Â  Â  Â  Â  Â  m_cost_hr = r * mining_cost
Â  Â  Â  Â  Â  Â  total_opex = labour_cost + p_cost + water_cost + maintenance_cost + m_cost_hr + lease_tax
Â  Â  Â  Â  Â  Â Â 
Â  Â  Â  Â  Â  Â  grid[i, j] = total_rev - total_opex
Â  Â  Â  Â  Â  Â Â 
Â  Â  return grid, rates, splits

# ----------- ADDITION: ECONOMICS CALCULATION -----------
power_cost = power_kw * power_rate
mining_cost_hr = f_rate * mining_cost

total_opex = (
Â  Â  labour_cost +
Â  Â  power_cost +
Â  Â  water_cost +
Â  Â  maintenance_cost +
Â  Â  mining_cost_hr +
Â  Â  lease_tax
)

profit_hr = total_revenue - total_opex
actual_margin = (profit_hr / total_revenue) * 100 if total_revenue > 0 else 0

kpi_status = {
Â  Â  "Throughput OK": f_rate >= target_throughput,
Â  Â "Profit Margin OK": actual_margin >= target_margin,
Â  Â  "Profit/hr OK": profit_hr >= target_profit_hr
}
cost_per_ton = total_opex / f_rate
profit_per_ton = profit_hr / f_rate

# --- 5. TABS ---
tab_viz, tab_theory, tab_export = st.tabs(["ğŸ“Š Dashboard", "ğŸ“– Theory & Equations", "ğŸ“¥ Export Report"])

with tab_viz:
Â  Â  c1, c2, c3 = st.columns(3)
Â  Â  c1.metric("Conc Flow", f"{c_mass:.2f} tph")
Â  Â  c2.metric("Total Revenue", f"${total_revenue:,.2f}/hr")
Â  Â  c3.metric("Feed Size", f"{d80} Âµm")

Â  Â  k1, k2, k3 = st.columns(3)
Â  Â  k1.metric("Throughput", f"{f_rate} tph")
Â  Â  k2.metric("Cost / ton", f"${cost_per_ton:.2f}")
Â  Â  k3.metric("Profit / hour", f"${profit_hr:,.2f}")

Â  Â  k4, k5 = st.columns(2)
Â  Â  k4.metric("Profit / ton", f"${profit_per_ton:.2f}")
Â  Â  k5.metric("Total OPEX", f"${total_opex:,.2f}/hr")

Â  Â  st.dataframe(df_res, use_container_width=True)

Â  Â  g1, g2 = st.columns(2)
Â  Â  with g1:
Â  Â  Â  Â  fig1, ax1 = plt.subplots()
Â  Â  Â  Â  ax1.pie(df_res["Revenue $/hr"], labels=df_res["Mineral"], autopct='%1.1f%%')
Â  Â  Â  Â  st.pyplot(fig1)

Â  Â  with g2:
Â  Â  Â  Â  fig2, ax2 = plt.subplots()
Â  Â  Â  Â  ax2.bar(df_res["Mineral"], df_res["Recovery %"])
Â  Â  Â  Â  ax2.set_ylabel("Recovery %")
Â  Â  Â  Â  st.pyplot(fig2)

Â  Â  fig3, ax3 = plt.subplots()
Â  Â  ax3.bar(["Revenue", "OPEX", "Profit"], [total_revenue, total_opex, profit_hr])
Â  Â  ax3.set_ylabel("$/hr")
Â  Â  st.pyplot(fig3)

with tab_theory:
Â  Â  st.subheader("ğŸ“š Engineering Logic")
Â  Â  st.latex(r"M_{feed} = M_{conc} + M_{midd} + M_{tail}")
Â  Â  st.latex(r"Mass\_Pull = 10\% + (Splitter\_Position \times 10\%)")
Â  Â  st.info(f"Operating at {f_rate} tph with splitter position {splitter_pos}")

with tab_export:
Â  Â  def make_report(df, rate, conc_m, d80_v, total_r, split_p):
Â  Â  Â  Â  doc = Document()
Â  Â  Â  Â  doc.add_heading('Engineering Report: Spiral Digital Twin', 0)
Â  Â  Â  Â  doc.add_paragraph(f"Feed: {rate} tph | d80: {d80_v} um | Splitter: {split_p}")
Â  Â  Â  Â Â 
Â  Â  Â  Â  fig_r, ax_r = plt.subplots()
Â  Â  Â  Â  ax_r.pie(df["Revenue $/hr"], labels=df["Mineral"], autopct='%1.1f%%')
Â  Â  Â  Â  img_s = BytesIO()
Â  Â  Â  Â  plt.savefig(img_s, format='png')
Â  Â  Â  Â  img_s.seek(0)
Â  Â  Â  Â  doc.add_picture(img_s, width=Inches(5))
Â  Â  Â  Â Â 
Â  Â  Â  Â  t = doc.add_table(rows=1, cols=len(df.columns))
Â  Â  Â  Â  t.style = 'Table Grid'
Â  Â  Â  Â  for i, col in enumerate(df.columns):
Â  Â  Â  Â  Â  Â  t.cell(0, i).text = col
Â  Â  Â  Â  for _, row in df.iterrows():
Â  Â  Â  Â  Â  Â  cells = t.add_row().cells
Â  Â  Â  Â  Â  Â  for i, val in enumerate(row):
Â  Â  Â  Â  Â  Â  Â  Â  cells[i].text = str(val)
Â  Â  Â  Â Â 
Â  Â  Â  Â  bio = BytesIO()
Â  Â  Â  Â  doc.save(bio)
Â  Â  Â  Â  return bio.getvalue()

st.write("---")
st.subheader("ğŸ”¥ Profitability Heatmap (Feed Rate vs Splitter)")

# Data generate karna
grid_data, x_labels, y_labels = generate_heatmap_data(user_targets, user_prices, d80, solids)

# Plotting
fig_heat, ax_heat = plt.subplots(figsize=(10, 6))
sns.heatmap(
Â  Â  grid_data,Â 
Â  Â  annot=False,Â 
Â  Â  xticklabels=np.round(y_labels, 1),Â 
Â  Â  yticklabels=np.round(x_labels, 0),Â 
Â  Â  cmap="RdYlGn",Â 
Â  Â  ax=ax_heat
)
ax_heat.set_xlabel("Splitter Position")
ax_heat.set_ylabel("Feed Rate (tph)")
st.pyplot(fig_heat)

st.info("ğŸ’¡ **Tip:** Green area sab se zyada profit dikhata hai. Red area ka matlab hai ke aapka OPEX revenue se zyada hai.")

Â  Â # --- Is niche wale block ko check karein aur "key" add karein ---
st.download_button(
Â  Â  label="ğŸ“¥ Download Full Executive Report (Word)",
Â  Â  data=make_report(df_res, f_rate, c_mass, d80, total_revenue, splitter_pos),
Â  Â  file_name="Spiral_Digital_Twin_Report.docx",
Â  Â  key="final_report_button",
) # <--- Ye bracket lazmi hai!

st.markdown("### ğŸ§  KPI Status Check")

for k, v in kpi_status.items():
Â  Â  if v:
Â  Â  Â  Â  st.success(f"âœ… {k}")
Â  Â  else:
Â  Â  Â  Â  st.error(f"âŒ {k}")

